"""
Read a procedure document from disk and return extracted plain text.

Supported:
- PDF
- DOCX
- TXT / MD



PDF flow:
1. Try normal text extraction using PyMuPDF.

2. If very little text is extracted, assume scanned PDF.
3. Fall back to OCR using pdf2image + Tesseract.
"""

from __future__ import annotations

from pathlib import Path
import os

from dotenv import load_dotenv

from audit.schemas.procedure_schema import ProcedureDocument


load_dotenv()


OCR_THRESHOLD = 100


def _parse_pages(pages_str: str) -> list[int]:
    """Parse a page string like '1, 3, 5-7' into a sorted list of 0-indexed page numbers."""
    result: set[int] = set()
    for part in pages_str.replace(" ", "").split(","):
        if not part:
            continue
        if "-" in part:
            a, _, b = part.partition("-")
            if a.isdigit() and b.isdigit():
                result.update(range(int(a) - 1, int(b)))  # 1-indexed → 0-indexed
        elif part.isdigit():
            result.add(int(part) - 1)
    return sorted(result)


def _read_pdf(file_path: Path, pages: list[int] | None = None) -> tuple[str, list[str], str]:
    warnings: list[str] = []

    try:
        import fitz
    except ImportError:
        return (
            "",
            ["PyMuPDF is not installed. Install it before testing PDF extraction."],
            "pdf_text",
        )

    document = fitz.open(file_path)
    total = len(document)
    page_indices = [i for i in (pages or range(total)) if 0 <= i < total]

    text = "\n".join(
        document[i].get_text("text")
        for i in page_indices
    ).strip()

    if len(text) >= OCR_THRESHOLD:
        return text, warnings, "pdf_text"

    warnings.append(
        "Very little text extracted. Assuming scanned PDF and running OCR."
    )

    try:
        from pdf2image import convert_from_path
        import pytesseract

        poppler_path = os.getenv("POPPLER_PATH")

        all_images = convert_from_path(str(file_path), poppler_path=poppler_path)
        selected_images = [all_images[i] for i in page_indices if i < len(all_images)]

        text = "\n".join(
            pytesseract.image_to_string(img) for img in selected_images
        ).strip()

        warnings.append("Scanned PDF detected. OCR used.")
        return text, warnings, "ocr"

    except ImportError:
        warnings.append("OCR dependencies missing. Install pdf2image and pytesseract.")
        return text, warnings, "pdf_text"

    except Exception as e:
        warnings.append(f"OCR failed: {e}")
        return text, warnings, "pdf_text"


def _read_docx(file_path: Path) -> tuple[str, list[str]]:
    try:
        from docx import Document

    except ImportError:
        return "", [
            "python-docx is not installed."
        ]

    document = Document(str(file_path))

    text = "\n".join(
        paragraph.text
        for paragraph in document.paragraphs
    ).strip()


    return text, []


def _read_text(file_path: Path) -> tuple[str, list[str]]:
    for encoding in (
        "utf-8",
        "latin-1",
        "utf-16"
    ):
        try:
            return (
                file_path.read_text(
                    encoding=encoding
                ).strip(),
                []
            )

        except UnicodeDecodeError:
            continue

    return "", [
        "Could not decode the text file."
    ]


def read_procedure_file(file_path: str, pages: list[int] | None = None) -> ProcedureDocument:
    path = Path(file_path)

    suffix = path.suffix.lower()

    source = "pdf_text"

    if suffix == ".pdf":
        text, warnings, source = _read_pdf(path, pages=pages)

    elif suffix == ".docx":
        text, warnings = _read_docx(path)

    elif suffix in {".txt", ".md"}:
        text, warnings = _read_text(path)

    else:
        text, warnings = "", [
            f"Unsupported procedure format: {suffix or 'unknown'}"
        ]

    # These warnings (OCR fallback, extraction failures, unreadable files) are
    # deliberately hidden from the UI as noise — print them so they're still
    # visible in the terminal when something silently produced empty/bad text.
    for w in warnings:
        print(f"  [procedure_reader] {path.name}: {w}")
    if not text.strip():
        print(f"  WARN: {path.name} produced no extractable text at all")

    return ProcedureDocument(
        file_path=str(path),
        file_type=suffix.lstrip("."),
        text=text,
        warnings=warnings,
        source=source,
    )